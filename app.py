import os
import openai
import streamlit as st
import docx
import time
import spacy


from dotenv import load_dotenv,find_dotenv
_ = load_dotenv(find_dotenv())
openai.api_key  = os.environ['OPEN_API_KEY_1']


def get_start(prompt, model="gpt-3.5-turbo"):
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0,
    )
    return response.choices[0].message["content"]


def document_process(input_file, language, name):
    nlp = spacy.load("en_core_web_lg")
    doc = docx.Document(input_file)

    batch_size = 5
    prompt_template = """
    Translate the following English text to '{}':

    '{}'
    """
    trans_res = []

    for batch_start in range(0, len(doc.paragraphs), batch_size):
        batch_end = min(batch_start + batch_size, len(doc.paragraphs))
        batch_paragraphs = doc.paragraphs[batch_start:batch_end]

        for paragraph in batch_paragraphs:
            text = paragraph.text
            doc = nlp(text)
            person_names = [ent.text for ent in doc.ents if ent.label_ == "PERSON"]

            txt = text
            for person_name in person_names:
                txt = txt.replace(person_name, name)

            prompt = prompt_template.format(language, txt)

            response = get_start(prompt)
            res = response.strip("'")
            trans_res.append(res)

            time.sleep(20)
    translated_text = "\n".join(trans_res)
    return translated_text


def main():
    st.header("Language Translater")
    input_file = st.file_uploader("Choose The File ")
    language = st.text_input("Enter Language")
    name = st.text_input("Enter Name")
    if st.button("Translate"):
        with st.spinner("Waiting for Translation..."):
            if input_file and language and name:
                translated_res = document_process(input_file, language, name)
                st.text_area(label="Translated text", value=translated_res)
                st.download_button('Download File', translated_res)


if __name__ == "__main__":
    main()